from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create document
doc = Document()
doc.add_heading('Budget and Budgetary Control - Notes', 0)

# Short Notes Section
doc.add_heading('🔹 SHORT NOTES (Quick Revision Format)', level=1)

short_notes = [
    ("Budget", "A financial plan for income & expenses over time (monthly/yearly). Helps manage resources & avoid overspending."),
    ("Key Concepts", "- Income: Money coming in (salary, sales, etc.)\n- Expenses: Money going out (bills, rent, etc.)\n- Savings: Income set aside for future\n- Surplus/Deficit: Income > Expenses (Surplus), Expenses > Income (Deficit)"),
    ("Importance of Budgeting", "- Ensures financial control\n- Aids goal setting\n- Helps manage debt\n- Prepares for emergencies"),
    ("Types of Budgets", "- Personal: For individuals/families\n- Business: For company finances\n- Government: For public revenue & spending"),
    ("Steps in Budgeting", "1. List income sources\n2. Track expenses (fixed & variable)\n3. Set financial goals\n4. Allocate for savings\n5. Monitor & adjust regularly")
]

for title, content in short_notes:
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)

# Fixed vs Flexible Budget Table
doc.add_heading("Fixed Budget vs. Flexible Budget", level=2)
table = doc.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Aspect'
hdr_cells[1].text = 'Fixed Budget'
hdr_cells[2].text = 'Flexible Budget'

rows = [
    ('Adaptability', 'Does not adjust', 'Adjusts with activity'),
    ('Complexity', 'Simple', 'Complex'),
    ('Best For', 'Stable environments', 'Dynamic environments'),
    ('Control', 'Strict', 'Realistic'),
    ('Example', '$10,000 fixed marketing spend', '5% of actual revenue for marketing')
]

for row in rows:
    cells = table.add_row().cells
    for i in range(3):
        cells[i].text = row[i]

# More Budgeting Types
doc.add_heading('More Budgeting Types', level=2)
types = {
    "Incremental Budgeting": "Builds on previous year's budget + adjustments\n✅ Simple, efficient\n❌ Can overlook inefficiencies",
    "Zero-Based Budgeting (ZBB)": "Starts from zero each cycle\n✅ Cost-efficient, goal-aligned\n❌ Time-consuming",
    "Activity-Based Budgeting (ABB)": "Based on specific business activities\nUseful in manufacturing, service sectors",
    "Performance-Based Budgeting (PBB)": "Links budget to outcomes/performance\nCommon in government and non-profits",
    "Rolling Budgeting": "Continuously updated by adding a new period (e.g., next month/quarter)"
}
for k, v in types.items():
    doc.add_heading(k, level=3)
    doc.add_paragraph(v)

# Brief Notes Section
doc.add_heading('🔹 BRIEF NOTES (Detailed Summary)', level=1)
brief_summary = (
    "Budget and Budgetary Control involve planning and controlling financial resources to meet short-term and long-term goals. "
    "Budgeting helps individuals, businesses, and governments allocate their income to various spending categories, maintain savings, and avoid deficits.\n\n"
    "Fixed Budgets are rigid and predetermined, ideal for predictable financial situations, offering simplicity and strong control. In contrast, Flexible Budgets adjust "
    "with real-time data like production levels, making them more suitable for variable environments.\n\n"
    "The Incremental Budgeting method uses last year’s budget as a base, making small adjustments. While quick and easy, it may perpetuate inefficiencies. "
    "Zero-Based Budgeting (ZBB) resets the budget every cycle, requiring all expenses to be justified from scratch—good for transparency but resource-intensive.\n\n"
    "Activity-Based Budgeting (ABB) emphasizes budgeting based on specific operations or activities, ideal for identifying cost drivers in production. "
    "Performance-Based Budgeting (PBB) ties funds to performance metrics, encouraging efficiency and results-based funding—often seen in the public sector.\n\n"
    "Finally, Rolling Budgets keep the financial plan dynamic by continuously adding new periods, making it a flexible approach to budgeting."
)
doc.add_paragraph(brief_summary)

# Save the document
file_path = r"D:\4thSem_MUJ\New folder"
doc.save(file_path)
file_path
